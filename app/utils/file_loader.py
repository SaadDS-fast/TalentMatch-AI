"""
File loading utilities for TalentMatch AI.

Supports plain text, PDF, and DOCX documents used by the API,
Streamlit frontend, parsers, and candidate ranking pipeline.
"""

from pathlib import Path


def load_text_from_file(file_path: str | Path) -> str:
    """
    Load text content from a supported document path.

    Args:
        file_path: Path to a .txt, .pdf, or .docx file.

    Returns:
        Extracted text.

    Raises:
        FileNotFoundError: If the path does not exist.
        ValueError: If the file extension is unsupported or text cannot be read.
    """

    path = Path(file_path).expanduser()

    if not path.exists():
        raise FileNotFoundError(f"File not found: {path}")

    if not path.is_file():
        raise ValueError(f"Expected a file path, got: {path}")

    suffix = path.suffix.lower()

    if suffix in {".txt", ".md"}:
        return _load_text(path)

    if suffix == ".pdf":
        return _load_pdf(path)

    if suffix == ".docx":
        return _load_docx(path)

    raise ValueError(
        f"Unsupported file type '{suffix}'. Supported types: .txt, .md, .pdf, .docx"
    )


def _load_text(path: Path) -> str:
    try:
        return path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        return path.read_text(encoding="latin-1")


def _load_pdf(path: Path) -> str:
    try:
        import pdfplumber

        pages = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                pages.append(page.extract_text() or "")

        text = "\n".join(page for page in pages if page.strip())
        if text.strip():
            return text
    except Exception:
        pass

    try:
        from PyPDF2 import PdfReader

        reader = PdfReader(str(path))
        pages = [page.extract_text() or "" for page in reader.pages]
        text = "\n".join(page for page in pages if page.strip())
        if text.strip():
            return text
    except Exception as exc:
        raise ValueError(f"Could not extract text from PDF: {path}") from exc

    raise ValueError(f"No extractable text found in PDF: {path}")


def _load_docx(path: Path) -> str:
    try:
        from docx import Document
    except Exception as exc:
        raise ValueError("python-docx is required to load DOCX files.") from exc

    try:
        document = Document(str(path))
        paragraphs = [paragraph.text for paragraph in document.paragraphs]
        table_cells = [
            cell.text
            for table in document.tables
            for row in table.rows
            for cell in row.cells
        ]
        return "\n".join([*paragraphs, *table_cells]).strip()
    except Exception as exc:
        raise ValueError(f"Could not extract text from DOCX: {path}") from exc
